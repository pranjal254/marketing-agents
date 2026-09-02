"""Deterministic Word (.docx) document builder — no LLM anywhere in this module.

Generic building blocks only: title, headings, paragraphs, key-value tables.
Domain templates (which sections a given document has) live in each agent package.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field


@dataclass(frozen=True)
class DocSection:
    heading: str
    paragraphs: tuple[str, ...] = ()
    table_rows: tuple[tuple[str, str], ...] = ()  # (label, value) pairs
    table_header: tuple[str, str] | None = None


@dataclass(frozen=True)
class DocSpec:
    title: str
    subtitle: str | None = None
    sections: tuple[DocSection, ...] = field(default_factory=tuple)


def build_docx(spec: DocSpec) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(spec.title, level=0)
    if spec.subtitle:
        doc.add_paragraph(spec.subtitle)
    for section in spec.sections:
        doc.add_heading(section.heading, level=1)
        for text in section.paragraphs:
            doc.add_paragraph(text)
        if section.table_rows:
            n_rows = len(section.table_rows) + (1 if section.table_header else 0)
            table = doc.add_table(rows=n_rows, cols=2)
            table.style = "Table Grid"
            offset = 0
            if section.table_header:
                table.cell(0, 0).text = section.table_header[0]
                table.cell(0, 1).text = section.table_header[1]
                offset = 1
            for i, (label, value) in enumerate(section.table_rows):
                table.cell(i + offset, 0).text = label
                table.cell(i + offset, 1).text = value
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
